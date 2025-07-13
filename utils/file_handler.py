"""
File upload and text extraction utilities, now with CSV/Excel support.
"""

import io
import pandas as pd
import PyPDF2
import mammoth
from docx import Document
from typing import Dict, Tuple, Optional


def process_uploaded_file(uploaded_file) -> str:
    """
    Process uploaded file and return content.
    Supports: txt, PDF, DOCX, CSV, XLS/XLSX.
    """
    try:
        file_content = ""
        mime = uploaded_file.type.lower()
        name = uploaded_file.name.lower()

        # Plain text
        if mime == "text/plain":
            file_content = uploaded_file.getvalue().decode("utf-8")

        # PDF
        elif mime == "application/pdf":
            file_content = extract_text_from_pdf(uploaded_file.getvalue())

        # DOCX
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            try:
                file_content = extract_text_from_docx(uploaded_file.getvalue())
            except Exception:
                file_content = extract_text_with_mammoth(uploaded_file.getvalue())

        # CSV
        elif mime == "text/csv" or name.endswith(".csv"):
            # Read CSV into DataFrame and convert to CSV text
            df = pd.read_csv(io.StringIO(uploaded_file.getvalue().decode("utf-8")))
            file_content = df.to_csv(index=False)

        # Excel (XLS or XLSX)
        elif mime in (
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ) or name.endswith(tuple([".xls", ".xlsx"])):
            # Read Excel into DataFrame and convert to CSV text
            df = pd.read_excel(io.BytesIO(uploaded_file.getvalue()))
            file_content = df.to_csv(index=False)

        else:
            # Fallback: decode bytes
            file_content = uploaded_file.getvalue().decode("utf-8", errors="ignore")

        return file_content.strip()
    except Exception as e:
        raise Exception(f"Error processing file: {str(e)}")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(
            [p.text for p in doc.paragraphs if p.text.strip()]
        )
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    [cell.text.strip() for cell in row.cells if cell.text.strip()]
                )
                if row_text:
                    text += "\n" + row_text
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def extract_text_with_mammoth(file_bytes: bytes) -> str:
    try:
        with io.BytesIO(file_bytes) as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            return result.value.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX with mammoth: {str(e)}")


def get_file_metadata(uploaded_file) -> Dict[str, any]:
    return {
        "name": uploaded_file.name,
        "type": uploaded_file.type,
        "size": uploaded_file.size,
        "key": f"{uploaded_file.name}_{uploaded_file.size}"
    }


def validate_file(
    uploaded_file,
    max_size: int,
    supported_types: list
) -> Tuple[bool, Optional[str]]:
    # Size check
    if uploaded_file.size > max_size:
        return False, (
            f"File size exceeds maximum of {max_size/1024/1024:.1f}MB"
        )
    # Extension check
    ext = uploaded_file.name.split('.')[-1].lower()
    if ext not in supported_types:
        return False, (
            f"Type '{ext}' not supported. Supported: {', '.join(supported_types)}"
        )
    return True, None
